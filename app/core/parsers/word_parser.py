"""
Word文件解析器
"""
import logging
from pathlib import Path
from typing import Dict, Any
from app.core.parsers.parser_interface import ParserInterface

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    logger.warning("python-docx未安装，Word解析功能不可用")


class WordParser(ParserInterface):
    """Word文件解析器（使用python-docx）"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Word文件为Markdown"""
        if not WORD_AVAILABLE:
            raise ImportError("python-docx未安装，无法解析Word文件。请运行: pip install python-docx")
        
        file = Path(file_path)
        
        if not file.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 读取Word文档
        doc = DocxDocument(file_path)
        
        # 提取所有段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 根据样式判断是否为标题
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        paragraphs.append(f"{'#' * level_num} {text}\n\n")
                    except ValueError:
                        paragraphs.append(f"## {text}\n\n")
                else:
                    paragraphs.append(f"{text}\n\n")
        
        content = "".join(paragraphs)
        
        # 提取标题（第一个段落或文件名）
        title = doc.paragraphs[0].text.strip()[:500] if doc.paragraphs and doc.paragraphs[0].text.strip() else file.stem[:500] if file.stem else "无标题"
        
        # 提取摘要（前几个段落的前200个字符）
        summary_text = " ".join([p.text.strip() for p in doc.paragraphs[:3] if p.text.strip()])
        summary = summary_text[:200] if summary_text else ""
        
        return {
            "content": content,
            "title": title,
            "summary": summary,
            "metadata": {
                "page_count": len(doc.paragraphs) // 20  # 粗略估算页数
            }
        }
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取Word文件元数据"""
        if not WORD_AVAILABLE:
            return {"page_count": None, "file_size": 0}
        
        file = Path(file_path)
        
        try:
            doc = DocxDocument(file_path)
            # 粗略估算页数（每页约20个段落）
            page_count = len(doc.paragraphs) // 20 if doc.paragraphs else 1
            
            return {
                "page_count": page_count,
                "file_size": file.stat().st_size if file.exists() else 0,
            }
        except Exception as e:
            logger.error(f"提取Word元数据失败: {e}")
            return {
                "page_count": None,
                "file_size": file.stat().st_size if file.exists() else 0,
            }

